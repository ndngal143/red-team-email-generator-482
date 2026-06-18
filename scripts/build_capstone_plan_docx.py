from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Capstone_Project_Plan.docx"
OLD_DRAFT_OUT = ROOT / "Capstone_Project_Plan_Draft.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def style_table(table, widths):
    set_table_width(table, widths)
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, heading in enumerate(headers):
        hdr[idx].text = heading
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
    style_table(table, widths)
    doc.add_paragraph()
    return table


def add_title_block(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Capstone Project Plan")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)
    run.bold = True
    title.paragraph_format.space_after = Pt(3)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Red Team Social Engineering Email Generator (For Authorized)")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    meta = [
        ("Course", "CSC482 Capstone Project II"),
        ("Team", "Team 3: Joyce, Sebastian, Alex"),
        ("Project Window", "June 15, 2026 through July 24, 2026"),
        ("Final Presentation/Demo", "July 27-28, 2026"),
    ]
    add_table(doc, ["Field", "Details"], meta, [1.8, 4.7])


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def strip_footers(path):
    tmp = path.with_suffix(".tmp.docx")
    footer_names = set()
    rel_footer_ids = set()

    with ZipFile(path, "r") as zin:
        document_xml = zin.read("word/document.xml")
        rels_xml = zin.read("word/_rels/document.xml.rels")

        rels_root = etree.fromstring(rels_xml)
        for rel in list(rels_root):
            if rel.get("Type", "").endswith("/footer"):
                rel_footer_ids.add(rel.get("Id"))
                footer_names.add("word/" + rel.get("Target", "").lstrip("/"))
                rels_root.remove(rel)

        doc_root = etree.fromstring(document_xml)
        for node in doc_root.xpath("//w:footerReference", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            parent = node.getparent()
            parent.remove(node)

        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in footer_names:
                    continue
                if item.filename == "word/document.xml":
                    zout.writestr(item, etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True))
                elif item.filename == "word/_rels/document.xml.rels":
                    zout.writestr(item, etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True))
                else:
                    data = zin.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        ct_root = etree.fromstring(data)
                        for override in list(ct_root):
                            part_name = override.get("PartName", "")
                            if part_name.lstrip("/") in footer_names:
                                ct_root.remove(override)
                        data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    zout.writestr(item, data)

    tmp.replace(path)


def build_doc():
    doc = Document()
    setup_styles(doc)
    add_title_block(doc)

    add_heading(doc, "Project Goal", 1)
    add_body(
        doc,
        "Build a web-based form generator for authorized red-team security awareness exercises. "
        "Users enter a target profile and scenario context, choose tone and deception intensity, "
        "and receive a customized phishing-style training email draft. The system will also include "
        "anti-spam indicator checks, realistic internal-looking training URL generation, and generated "
        "email logs for defensive model training.",
    )

    add_heading(doc, "Safety and Scope", 1)
    add_body(
        doc,
        "This project is for authorized internal awareness exercises only. Generated links should be "
        "training/demo links, not credential-harvesting pages. Logs should support defensive analysis "
        "and model training, not real-world abuse.",
    )

    add_heading(doc, "Timeline Overview", 1)
    add_table(
        doc,
        ["Dates", "Milestone"],
        [
            ("6/15-6/19", "Milestone 1: Requirements, Architecture, and Project Setup"),
            ("6/22-6/26", "Milestone 2: Form Design, Data Model, and Safety Controls"),
            ("6/29-7/3", "Milestone 3: Email Generation Engine and Spam Indicator Checks"),
            ("7/6-7/10", "Milestone 4: Link Generator, Logging, and Storage"),
            ("7/13-7/17", "Milestone 5: Full Integration, Testing, and Defensive Dataset Export"),
            ("7/20-7/24", "Milestone 6: Polish, Deployment, Documentation, and Demo Prep"),
            ("7/27-7/28", "Final Presentation and Demonstration"),
        ],
        [1.3, 5.2],
    )

    milestones = [
        (
            "Milestone 1: Requirements, Architecture, and Project Setup",
            "6/15-6/19",
            "Define the project clearly, set up the development environment, and create the technical foundation.",
            [
                ("Joyce", "Define user requirements and project use cases", "Requirements document, user stories, feature list", "At least 8 user stories; requirements reviewed by all 3 members; scope includes generator, tone controls, intensity controls, link generator, and logging"),
                ("Sebastian", "Design system architecture", "Architecture diagram, technology stack decision, data flow diagram", "Architecture includes frontend, backend/API, LLM generation layer, URL generator, logging storage, and safety checks"),
                ("Alex", "Set up repository and development environment", "GitHub repository, initial folder structure, README, setup instructions", "Repo runs locally; README includes install/run steps; all team members can clone and start the app"),
            ],
        ),
        (
            "Milestone 2: Form Design, Data Model, and Safety Controls",
            "6/22-6/26",
            "Build the user input workflow and define the structured data needed for generation and logging.",
            [
                ("Joyce", "Design web form fields and user workflow", "Form wireframe, field list, validation rules", "Form includes target profile, department, scenario context, tone, deception intensity, sender role, and organization details"),
                ("Sebastian", "Create backend data model", "Schema for generation requests, generated emails, generated links, and logs", "Schema supports all required form fields; log entries include timestamp, inputs, generated output, spam score, and link metadata"),
                ("Alex", "Define safety and authorization controls", "Safety checklist, warning labels, allowed-use notice, blocked content rules", "App displays authorized-use notice; generated content avoids credential collection instructions and real malicious payloads"),
            ],
        ),
        (
            "Milestone 3: Email Generation Engine and Spam Indicator Checks",
            "6/29-7/3",
            "Implement the core email draft generator and basic anti-spam analysis.",
            [
                ("Joyce", "Create tone and intensity prompt templates", "Prompt template set for urgent, professional, and friendly tones", "Each tone produces clearly different wording; deception intensity changes subject line, urgency, and call-to-action strength"),
                ("Sebastian", "Implement backend generation endpoint", "API endpoint that accepts form data and returns generated email draft", "Endpoint returns subject, sender name, body, call-to-action, and generated link placeholder within 5 seconds locally"),
                ("Alex", "Build spam indicator checker", "Spam keyword/risk scoring function and recommendations", "Checker flags common spam indicators such as excessive urgency, suspicious wording, all-caps, too many exclamation points, and risky phrases"),
            ],
        ),
        (
            "Milestone 4: Link Generator, Logging, and Storage",
            "7/6-7/10",
            "Add realistic internal-looking training URLs and persistent logs for defensive analysis.",
            [
                ("Joyce", "Define internal-looking URL patterns", "URL pattern list and sample generated links", "At least 10 realistic training-safe URL formats, such as policy, HR, finance, IT, and benefits paths"),
                ("Sebastian", "Implement link generator", "Function/API that generates training URLs based on department and scenario", "Generated links contain safe demo domains or local routes; no real credential capture URLs are produced"),
                ("Alex", "Implement email logging", "Log storage, generated email history view, exportable records", "Every generation creates a log entry; logs can be viewed and exported as CSV or JSON"),
            ],
        ),
        (
            "Milestone 5: Full Integration, Testing, and Defensive Dataset Export",
            "7/13-7/17",
            "Connect all components, test the full workflow, and prepare logs for defensive model training.",
            [
                ("Joyce", "Conduct user workflow testing", "Test cases, user feedback notes, revised UI checklist", "At least 10 test scenarios covering different departments, tones, and intensity levels"),
                ("Sebastian", "Integrate frontend, backend, generator, link generator, and logging", "Working end-to-end application", "User can submit form, generate email, receive spam check results, get generated URL, and see log entry"),
                ("Alex", "Create defensive dataset export", "CSV/JSON export format with labeled generated emails", "Export includes email body, subject, tone, intensity, scenario, spam indicators, and generated link metadata"),
            ],
        ),
        (
            "Milestone 6: Polish, Deployment, Documentation, and Demo Prep",
            "7/20-7/24",
            "Prepare the final working project for demonstration and presentation.",
            [
                ("Joyce", "Polish UI and prepare demo script", "Final UI revisions, demo walkthrough script", "Demo script covers form input, email generation, spam check, URL generation, and log export"),
                ("Sebastian", "Prepare deployment or local demo environment", "Deployed app or reliable local demo setup", "App runs consistently on demo machine; setup time under 5 minutes"),
                ("Alex", "Write final documentation and presentation content", "Final README, project report sections, presentation slides outline", "Documentation explains purpose, setup, features, safety limits, architecture, and testing results"),
            ],
        ),
    ]

    for title, dates, goal, rows in milestones:
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_heading(doc, title, 1)
        add_body(doc, f"Dates: {dates}")
        add_body(doc, f"Goal: {goal}")
        add_table(doc, ["Member", "Subtask", "Outputs", "Measurement"], rows, [1.0, 1.8, 1.8, 1.9])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Final Presentation and Demonstration", 1)
    add_body(doc, "Dates: 7/27-7/28")
    add_heading(doc, "Presentation Focus", 2)
    add_bullets(
        doc,
        [
            "Problem statement: why organizations need authorized phishing-awareness training.",
            "Project goal and scope.",
            "System architecture.",
            "Live demonstration of the web form generator.",
            "Demonstration of tone and deception intensity controls.",
            "Demonstration of anti-spam indicator checks.",
            "Demonstration of internal-looking training URL generation.",
            "Demonstration of generated email logs and defensive export.",
            "Lessons learned and future improvements.",
        ],
    )

    add_heading(doc, "Final Demo Success Criteria", 2)
    add_table(
        doc,
        ["Demo Area", "Measurement"],
        [
            ("Web form", "All required fields accept and validate input"),
            ("Email generation", "Produces complete email with subject, sender, body, and call-to-action"),
            ("Tone control", "Urgent, professional, and friendly outputs are visibly different"),
            ("Intensity control", "Low, medium, and high settings affect urgency and persuasion level"),
            ("Spam checker", "Flags common risky indicators and provides a score/recommendations"),
            ("Link generator", "Produces realistic but safe internal-looking training URLs"),
            ("Logging", "Each generated email is saved with metadata"),
            ("Export", "Logs export successfully as CSV or JSON"),
            ("Demo reliability", "Full workflow completes without errors during presentation"),
        ],
        [1.7, 4.8],
    )

    add_heading(doc, "Suggested Weekly Meeting Schedule", 1)
    add_table(
        doc,
        ["Date", "Meeting Focus"],
        [
            ("6/15", "Kickoff, assign roles, confirm tools"),
            ("6/19", "Review requirements and architecture"),
            ("6/26", "Review form design, data model, and safety rules"),
            ("7/3", "Review generation engine and spam checker"),
            ("7/10", "Review link generator and logs"),
            ("7/17", "End-to-end testing review"),
            ("7/24", "Final demo rehearsal"),
            ("7/27-7/28", "Final presentation and demonstration"),
        ],
        [1.3, 5.2],
    )

    doc.save(OUT)
    strip_footers(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
